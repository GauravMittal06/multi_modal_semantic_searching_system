# src/extractor.py
"""
Document Structure Extractor — Module 1
Extracts headings, paragraphs, tables, and images from PDF and DOCX files.
Outputs every element as a structured dict matching the BRD JSON schema.
"""
import re
import os
import uuid
import json
from typing import List, Dict, Any

import fitz  # PyMuPDF


# ─── JSON Schema helpers ─────────────────────────────────────────────────────

def _make_element(
    elem_type: str,
    content: str,
    page_number: int,
    section_heading: str,
    source_document: str,
    extra: Dict = None,
) -> Dict[str, Any]:
    base = {
        "id": str(uuid.uuid4()),
        "type": elem_type,
        "content": content,
        "page_number": page_number,
        "section_heading": section_heading,
        "source_document": source_document,
        "related_elements": [],
        "metadata": extra or {},
    }
    return base


# ─── PDF Extractor ────────────────────────────────────────────────────────────
def _line_text(line):
    spans = line.get("spans", [])
    return " ".join(
        s.get("text", "").strip()
        for s in spans
        if s.get("text", "").strip()
    ).strip()

def extract_from_pdf(filepath: str, source_name: str) -> List[Dict[str, Any]]:
    elements: List[Dict[str, Any]] = []
    current_heading = ""

    # --- Text + Headings via PyMuPDF ---
    doc = fitz.open(filepath)
    _all_doc_sizes = [
        span.get("size", 12)
        for page in doc
        for block in page.get_text("dict")["blocks"]
        if block.get("type") == 0
        for line in block.get("lines", [])
        for span in line.get("spans", [])
        if span.get("text", "").strip()
    ]
    doc_median_font = sorted(_all_doc_sizes)[len(_all_doc_sizes) // 2] if _all_doc_sizes else 12

    for page_num, page in enumerate(doc, start=1):
        blocks = page.get_text("dict")["blocks"]
        all_sizes = [span.get("size", 12) for block in blocks if block.get("type") == 0
                     for line in block.get("lines", []) for span in line.get("spans", [])
                     if span.get("text", "").strip()]
        page_median_font = max(
            sorted(all_sizes)[len(all_sizes) // 2] if all_sizes else 12,
            doc_median_font,
        )
        _page_heading_start_idx = len(elements)
        for block in blocks:
            if block.get("type") != 0:
                continue
            
            block_lines = []
            max_font_size = 0
            has_bold = False

            for line in block.get("lines", []):
                text = _line_text(line)

                if not text:
                    continue
                
                block_lines.append(text)

                spans = line.get("spans", [])

                line_font = max(
                    (s.get("size", 12) for s in spans),
                    default=12
                )

                max_font_size = max(max_font_size, line_font)

                if any(bool(s.get("flags", 0) & (2**4)) for s in spans):
                    has_bold = True

            block_text = " ".join(block_lines).strip()

            if len(block_text) < 5:
                continue
            
            words = block_text.split()

            digit_ratio = (
                sum(
                    1
                    for w in words
                    if w.replace("$", "")
                    .replace("%", "")
                    .replace(",", "")
                    .replace(".", "")
                    .replace("+", "")
                    .replace("-", "")
                    .isdigit()
                )
                / max(len(words), 1)
            )

            is_heading = (
                (max_font_size > page_median_font * 1.15
                 or (has_bold and max_font_size > page_median_font * 1.05))
                and len(block_text) < 120
                and 2 <= len(words) <= 20
                and digit_ratio < 0.15
                and not block_text.endswith((".", ":", ";", ",", "?"))
            )

            _HEADING_BLOCKLIST = {"n", "o", "no", "of", "to", "by", "in", "or", "and"}
            _SIDEBAR_PATTERNS = {
                "speak up", "reporting concerns", "speak up | reporting concerns",
                "see accompanying notes", "additional restrictions", "designated individuals",
            }
            if (
                block_text.lower() in _HEADING_BLOCKLIST
                or block_text.lower() in _SIDEBAR_PATTERNS
                or any(p in block_text.lower() for p in _SIDEBAR_PATTERNS)
                or block_text.strip() in {"}", "{", ")", "(", "]", "[", "|", "—", "-", "–"}
                or re.match(r"^\d+[\.\)\-]?\s*$", block_text.strip())
                or all(len(w) == 1 for w in words)
            ):
                is_heading = False

            if is_heading:
                current_heading = block_text
                _bbox = block.get("bbox", [0, 0, 0, 0])
                block_y0 = _bbox[1]

                elements.append(
                    _make_element(
                        "heading",
                        block_text,
                        page_num,
                        current_heading,
                        source_name,
                        {"font_size": max_font_size, "y0": block_y0, "bbox": list(_bbox)},
                    )
                )

                # ===== PART_B_BBOX_PROBE =====
                print(
                    f"[PART_B_BBOX_PROBE] page={page_num} | "
                    f"bbox=({_bbox[0]:.2f}, {_bbox[1]:.2f}, {_bbox[2]:.2f}, {_bbox[3]:.2f}) | "
                    f"font_size={max_font_size:.2f} | "
                    f"text={block_text!r}"
                )

            else:
                if len(block_text.strip()) >= 5:
                    elements.append(
                        _make_element(
                            "paragraph",
                            block_text,
                            page_num,
                            current_heading,
                            source_name,
                            {},
                        )
                    )

        # ===== PART_B: multiline heading reconstruction (this page only) =====
        _FONT_TOL = 0.5
        _GAP_MIN, _GAP_MAX = -5, 5
        _TERMINAL_PUNCT = (".", ":", ";", ",", "?", "!")

        _page_elements = elements[_page_heading_start_idx:]
        _i = 0
        while _i < len(_page_elements) - 1:
            a = _page_elements[_i]
            b = _page_elements[_i + 1]

            if a["type"] != "heading" or b["type"] != "heading":
                _i += 1
                continue

            a_bbox = a.get("metadata", {}).get("bbox")
            b_bbox = b.get("metadata", {}).get("bbox")
            a_fs = a.get("metadata", {}).get("font_size")
            b_fs = b.get("metadata", {}).get("font_size")

            if not a_bbox or not b_bbox or a_fs is None or b_fs is None:
                _i += 1
                continue

            same_font = abs(a_fs - b_fs) <= _FONT_TOL
            gap = b_bbox[1] - a_bbox[3]  # b.y0 - a.y1
            gap_ok = _GAP_MIN <= gap <= _GAP_MAX
            indented = b_bbox[0] > a_bbox[0]
            no_term_punct = not a["content"].rstrip().endswith(_TERMINAL_PUNCT)

            if same_font and gap_ok and indented and no_term_punct:
                merged_text = f"{a['content']} {b['content']}"
                print(f'[PART_B_MERGE] MERGED: "{a["content"]}" + "{b["content"]}" -> "{merged_text}"')

                a["content"] = merged_text
                a["section_heading"] = merged_text
                a["metadata"]["bbox"] = [
                    min(a_bbox[0], b_bbox[0]),
                    a_bbox[1],
                    max(a_bbox[2], b_bbox[2]),
                    b_bbox[3],
                ]
                # y0 stays as a's original top position

                # Remove b from the real elements list (it's a duplicate now)
                elements.remove(b)
                _page_elements.pop(_i + 1)

                # If current_heading was set to b's standalone text, fix it forward
                if current_heading == b["content"]:
                    current_heading = merged_text

                # Re-check the new a against the next element (don't advance _i)
                continue
            else:
                if a["type"] == "heading" and b["type"] == "heading":
                    print(
                        f'[PART_B_MERGE] SKIPPED: "{a["content"]}" + "{b["content"]}" '
                        f'| reason: '
                        f'{"font mismatch " if not same_font else ""}'
                        f'{"gap out of range (%.2f) " % gap if not gap_ok else ""}'
                        f'{"not indented " if not indented else ""}'
                        f'{"terminal punctuation " if not no_term_punct else ""}'
                    )
                _i += 1
        # ===== END PART_B =====

    # --- Tables via pdfplumber ---
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                tables = page.find_tables()
                for table in tables:
                    table = table.extract() if hasattr(table, "extract") else table
                    if not table:
                        continue
                    # Convert table rows to markdown-style string
                    rows = []
                    for row in table:
                        cleaned = [str(cell).strip() if cell else "" for cell in row]
                        rows.append(" | ".join(cleaned))
                    table_text = "\n".join(rows)
                    if table_text.strip():
                        # Prepend column headers as a hint so the LLM knows what the table contains
                        header_hint = " | ".join([str(c).strip() for c in table[0]]) if table else ""
                        table_text_clean = table_text.replace("$", "USD ").replace("%", " pct")
                        header_hint_clean = header_hint.replace("$", "USD ").replace("%", " pct")
                        all_cells = []
                        for row in table[1:]:
                            for cell in row:
                                if cell and len(str(cell).strip()) > 2:
                                    all_cells.append(str(cell).strip().replace("$", "USD ").replace("%", " pct"))
                        semantic_tags = " | ".join(dict.fromkeys(all_cells[:12]))
                        table_text_with_hint = f"[TABLE COLUMNS: {header_hint_clean}]\n[KEY VALUES: {semantic_tags}]\n{table_text_clean}"
                        _table_bbox = page.find_tables().tables[0].bbox if False else None  # placeholder removed below
                        elem = _make_element(
                            "table", table_text_with_hint, page_num,
                            _get_heading_for_page(elements, page_num),
                            source_name,
                            {
                                "raw_rows": len(table),
                                "raw_cols": len(table[0]) if table else 0,
                                "table_header": header_hint_clean,
                            }
                        )
                        elem["embed_content"] = table_text_clean
                        elements.append(elem)
    except ImportError:
        print("[extractor] pdfplumber not installed — skipping table extraction")
    except Exception as e:
        print(f"[extractor] table extraction error: {e}")

    # --- Images via PyMuPDF (pixmap per-page for vector charts) ---
    for page_num, page in enumerate(doc, start=1):
        images_dir = os.path.join(os.path.dirname(filepath), "extracted_images")
        os.makedirs(images_dir, exist_ok=True)

        # First try embedded raster images
        image_list = page.get_images(full=True)
        for img_index, img in enumerate(image_list):
            xref = img[0]
            try:
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                image_filename = f"{source_name}_p{page_num}_img{img_index}.{image_ext}"
                image_path = os.path.join(images_dir, image_filename)
                with open(image_path, "wb") as f:
                    f.write(image_bytes)
                _img_y0 = img[7] if len(img) > 7 else None
                # img tuple format from page.get_images(full=True) does not reliably
                # include placement bbox; fetch actual placement rect explicitly.
                try:
                    _img_rects = page.get_image_rects(xref)
                    _img_y0 = _img_rects[0].y0 if _img_rects else None
                except Exception:
                    _img_y0 = None

                elem = {
                    "id": str(uuid.uuid4()),
                    "type": "image",
                    "content": "",
                    "page_number": page_num,
                    "section_heading": _get_heading_for_page(elements, page_num, target_y0=_img_y0),
                    "source_document": source_name,
                    "related_elements": [],
                    "metadata": {"image_path": image_path, "image_filename": image_filename, "extension": image_ext, "y0": _img_y0},
                    "vision_summary": "",
                    "keywords": [],
                }
                elements.append(elem)
            except Exception as e:
                print(f"[extractor] raster image error page {page_num} img {img_index}: {e}")

        # Render full page as pixmap to capture vector charts/figures
        try:
            drawings = page.get_drawings()
            page_rect = page.rect
            total_drawings = len(drawings)

            ZONE_COUNT = 6
            zone_height = page_rect.height / ZONE_COUNT
            zone_drawing_counts = [0] * ZONE_COUNT

            for d in drawings:
                r = d.get("rect")
                if not r:
                    continue
                r = fitz.Rect(r)
                # Exclude full-width elements — page borders, horizontal rules, table edges
                if r.width > page_rect.width * 0.80:
                    continue
                center_y = (r.y0 + r.y1) / 2
                zone_idx = min(int(center_y / zone_height), ZONE_COUNT - 1)
                zone_drawing_counts[zone_idx] += 1

            max_zone_density = max(zone_drawing_counts) if zone_drawing_counts else 0

            # ── Hybrid trigger: local clustering AND overall complexity ──────
            ZONE_DENSITY_THRESHOLD = 7
            TOTAL_DRAWINGS_THRESHOLD = 15

            has_vector_graphics = (
                max_zone_density >= ZONE_DENSITY_THRESHOLD
                and total_drawings >= TOTAL_DRAWINGS_THRESHOLD
            )

            print(
                f"[extractor] Page {page_num}: total_drawings={total_drawings} | "
                f"zones={zone_drawing_counts} | max_zone_density={max_zone_density} | "
                f"trigger={'RENDER' if has_vector_graphics else 'skip'}"
            )

            # Removed 'and not image_list' so charts are always captured
            if has_vector_graphics:
                mat = fitz.Matrix(2.0, 2.0)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                page_image_filename = f"{source_name}_p{page_num}_fullpage.png"
                page_image_path = os.path.join(images_dir, page_image_filename)
                pix.save(page_image_path)
                elem = {
                    "id": str(uuid.uuid4()),
                    "type": "image",
                    "content": "",
                    "page_number": page_num,
                    "section_heading": _get_heading_for_page(elements, page_num),
                    "source_document": source_name,
                    "related_elements": [],
                    "metadata": {
                        "image_path": page_image_path,
                        "image_filename": page_image_filename,
                        "extension": "png",
                        "is_page_render": True,
                        "image_type": "page_snapshot",
                        "max_zone_density": max_zone_density,
                        "total_drawings": total_drawings,
                    },
                    "vision_summary": "",
                    "keywords": [],
                }
                elements.append(elem)

        except Exception as e:
            print(f"[extractor] vector/page render error page {page_num}: {e}")

    print("\n=== HEADING DEBUG ===")

    for e in elements:
        if e["type"] == "heading":
            print(
                f"page={e['page_number']} | "
                f"heading={e['content']}"
            )

    print("=====================\n")

    print("\n=== IMAGE SECTION ASSIGNMENT DEBUG (Part A validation) ===")

    for e in elements:
        if e["type"] != "image":
            continue

        page_num = e["page_number"]
        img_y0 = e.get("metadata", {}).get("y0")

        same_page_headings = [
            h for h in elements
            if h["type"] == "heading" and h["page_number"] == page_num
        ]

        print(
            f"\nimage page={page_num} | "
            f"assigned_section={e.get('section_heading')!r} | "
            f"image_y0={img_y0}"
        )
        if not same_page_headings:
            print(f"  (no headings on page {page_num})")
        else:
            for h in same_page_headings:
                h_y0 = h.get("metadata", {}).get("y0")
                above = (
                    "ABOVE image" if (h_y0 is not None and img_y0 is not None and h_y0 <= img_y0)
                    else "BELOW image" if (h_y0 is not None and img_y0 is not None)
                    else "y0 unknown"
                )
                marker = " <-- ASSIGNED" if h["content"] == e.get("section_heading") else ""
                print(f"  heading y0={h_y0} | {above} | {h['content']!r}{marker}")

    print("=== END IMAGE SECTION ASSIGNMENT DEBUG ===\n")

    doc.close()
    return elements


# ─── DOCX Extractor ───────────────────────────────────────────────────────────

def extract_from_docx(filepath: str, source_name: str) -> List[Dict[str, Any]]:
    from docx import Document
    from docx.oxml.ns import qn

    elements: List[Dict[str, Any]] = []
    doc = Document(filepath)
    current_heading = ""
    page_num = 1

    images_dir = os.path.join(os.path.dirname(filepath), "extracted_images")
    os.makedirs(images_dir, exist_ok=True)
    _docx_img_counter = 0

    def _extract_images_from_paragraph(paragraph, page_num, current_heading):
        nonlocal _docx_img_counter
        imgs = []
        blips = paragraph._p.findall(
            './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
        )
        for blip in blips:
            rid = blip.get(qn('r:embed'))
            if not rid:
                continue
            try:
                image_part = doc.part.related_parts[rid]
                image_bytes = image_part.blob
                image_ext = image_part.content_type.split("/")[-1]
                if image_ext in ("x-emf", "x-wmf"):
                    continue  # vector/metafile formats, not raster — skip
                image_filename = f"{source_name}_p{page_num}_img{_docx_img_counter}.{image_ext}"
                image_path = os.path.join(images_dir, image_filename)
                with open(image_path, "wb") as f:
                    f.write(image_bytes)
                _docx_img_counter += 1
                imgs.append({
                    "id": str(uuid.uuid4()),
                    "type": "image",
                    "content": "",
                    "page_number": page_num,
                    "section_heading": current_heading,
                    "source_document": source_name,
                    "related_elements": [],
                    "metadata": {
                        "image_path": image_path,
                        "image_filename": image_filename,
                        "extension": image_ext,
                        "y0": None,
                    },
                    "vision_summary": "",
                    "keywords": [],
                })
            except Exception as e:
                print(f"[extractor] docx image extract error page {page_num}: {e}")
        return imgs

    all_font_sizes = [r.font.size.pt for p in doc.paragraphs for r in p.runs if r.font.size]
    doc_median_font = sorted(all_font_sizes)[len(all_font_sizes)//2] if all_font_sizes else 11

    # crude page estimator — docx has no real page numbers, so we
    # accumulate character count and roll over every ~CHARS_PER_PAGE
    CHARS_PER_PAGE = 3000
    char_count = 0

    def _has_page_break(paragraph) -> bool:
        brs = paragraph._p.xpath('.//w:br[@w:type="page"]')
        return len(brs) > 0

    def _is_heading_like(paragraph) -> bool:
        style = paragraph.style.name or ""
        if style.startswith("Heading") or style.startswith("Title"):
            return True
        text = paragraph.text.strip()
        if not text or len(text) > 90:
            return False
        words = text.split()
        if len(text.strip()) < 8:
            return False
        if len(words) < 2:
            return False
        # fallback heuristic: bold/large-font short paragraphs act as headings
        # even when not tagged with a Word "Heading" style
        runs = paragraph.runs
        if not runs:
            return False
        bold_run_chars = sum(len(r.text) for r in runs if r.bold)
        total_chars = sum(len(r.text) for r in runs) or 1
        mostly_bold = (bold_run_chars / total_chars) > 0.6
        large_font = any((r.font.size and r.font.size.pt > doc_median_font * 1.15) for r in runs)
        is_upper_or_titled = text.isupper() or text.istitle()
        return mostly_bold and (large_font or is_upper_or_titled)

    # build a table -> preceding-paragraph-index map so tables can be
    # inserted at the correct position relative to surrounding text
    body_children = list(doc.element.body)
    table_xml_elements = {id(t._tbl): t for t in doc.tables}

    para_iter = iter(doc.paragraphs)

    for child in body_children:
        tag = child.tag.split('}')[-1]

        if tag == "p":
            para = next(para_iter, None)
            if para is None:
                continue
            text = para.text.strip()

            if _has_page_break(para):
                page_num += 1
                char_count = 0

            if not text:
                img_elems = _extract_images_from_paragraph(para, page_num, current_heading)
                elements.extend(img_elems)
                continue

            char_count += len(text)
            if char_count > CHARS_PER_PAGE:
                page_num += 1
                char_count = 0

            if _is_heading_like(para):
                current_heading = text
                elements.append(_make_element(
                    "heading", text, page_num, current_heading,
                    source_name, {"style": para.style.name or ""}
                ))
            elif len(text) > 20:
                elements.append(_make_element(
                    "paragraph", text, page_num, current_heading,
                    source_name, {}
                ))

            img_elems = _extract_images_from_paragraph(para, page_num, current_heading)
            elements.extend(img_elems)

        elif tag == "tbl":
            tbl_obj = table_xml_elements.get(id(child))
            if tbl_obj is None:
                continue
            rows = []
            for row in tbl_obj.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_text = "\n".join(rows)
            if table_text.strip():
                table_text_clean = table_text.replace("$", "USD ").replace("%", " pct")
                header_hint = rows[0] if rows else ""
                all_cells = []
                for row in tbl_obj.rows[1:]:
                    for cell in row.cells:
                        ct = cell.text.strip().replace("$", "USD ").replace("%", " pct")
                        if ct and len(ct) > 2:
                            all_cells.append(ct)
                semantic_tags = " | ".join(dict.fromkeys(all_cells[:12]))
                table_text_with_hint = f"[TABLE COLUMNS: {header_hint}]\n[KEY VALUES: {semantic_tags}]\n{table_text_clean}"
                tbl_elem = _make_element(
                    "table", table_text_with_hint, page_num, current_heading,
                    source_name, {
                        "raw_rows": len(tbl_obj.rows),
                        "raw_cols": len(tbl_obj.columns),
                        "table_header": header_hint,
                    }
                )
                tbl_elem["embed_content"] = table_text_clean
                elements.append(tbl_elem)
                char_count += len(table_text)

            for row in tbl_obj.rows:
                for cell in row.cells:
                    for cell_para in cell.paragraphs:
                        img_elems = _extract_images_from_paragraph(cell_para, page_num, current_heading)
                        elements.extend(img_elems)

    return elements


# ─── Universal Entry Point ────────────────────────────────────────────────────

def extract_document(filepath: str, source_name: str = None) -> List[Dict[str, Any]]:
    if not source_name:
        source_name = os.path.basename(filepath)
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return extract_from_pdf(filepath, source_name)
    elif ext in (".docx", ".doc"):
        return extract_from_docx(filepath, source_name)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: .pdf, .docx")
    
# ─── Internal helpers ─────────────────────────────────────────────────────────

def _get_heading_for_page(elements: List[Dict], page_num: int, target_y0: float = None) -> str:
    """
    Prefer the nearest heading ABOVE this element on the same page, by y-position.
    Falls back to the last heading on the page if no y-position is available
    (preserves old behavior for any caller that doesn't pass target_y0).
    Otherwise only inherit a heading from a nearby page.
    """

    MAX_HEADING_LOOKBACK = 3

    page_headings = [
        e for e in elements
        if e["type"] == "heading"
        and e["page_number"] == page_num
    ]

    if page_headings:
        if target_y0 is not None:
            # Headings with a known y0 that sit above (or at) the target element,
            # i.e. they were rendered earlier vertically on the page.
            candidates = [
                h for h in page_headings
                if h.get("metadata", {}).get("y0") is not None
                and h["metadata"]["y0"] <= target_y0
            ]
            if candidates:
                # Nearest above = largest y0 among those still <= target_y0
                nearest = max(candidates, key=lambda h: h["metadata"]["y0"])
                return nearest["content"]
            # No heading above this element on the page (e.g. it's above all
            # headings) — fall through to extraction-order-last as a safe default,
            # same as old behavior, rather than silently returning "".
            return page_headings[-1]["content"]

        # No target_y0 provided by caller — preserve exact old behavior.
        return page_headings[-1]["content"]

    closest_heading = ""
    closest_page = None

    for e in elements:
        if e["type"] != "heading":
            continue

        if e["page_number"] >= page_num:
            continue

        if (
            closest_page is None
            or e["page_number"] > closest_page
        ):
            closest_page = e["page_number"]
            closest_heading = e["content"]

    if (
        closest_page is not None
        and page_num - closest_page <= MAX_HEADING_LOOKBACK
    ):
        return closest_heading

    return ""